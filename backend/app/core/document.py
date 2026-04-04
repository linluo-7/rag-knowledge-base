"""
文档解析模块
支持解析 Word (.docx) 等格式的文档
"""

import re
from pathlib import Path
from typing import Optional
from docx import Document as DocxDocument


class DocumentParser:
    """
    文档解析器
    
    支持格式：
    - .docx (Word 文档)
    - .txt (纯文本)
    """
    
    def __init__(self):
        self.supported_formats = [".docx", ".txt"]
    
    def parse(self, file_path: str) -> str:
        """
        解析文档，返回纯文本内容
        
        Args:
            file_path: 文件路径
            
        Returns:
            str: 文档的纯文本内容
        """
        path = Path(file_path)
        suffix = path.suffix.lower()
        
        if suffix not in self.supported_formats:
            raise ValueError(f"不支持的文件格式: {suffix}，支持的格式: {self.supported_formats}")
        
        if suffix == ".docx":
            return self._parse_docx(path)
        elif suffix == ".txt":
            return self._parse_txt(path)
    
    def _parse_docx(self, path: Path) -> str:
        """
        解析 Word 文档
        
        使用 python-docx 库读取文档内容，
        提取所有段落文本并合并。
        """
        doc = DocxDocument(path)
        paragraphs = []
        
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                paragraphs.append(text)
        
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_text.append(cell_text)
                if row_text:
                    paragraphs.append(" | ".join(row_text))
        
        return "\n".join(paragraphs)
    
    def _parse_txt(self, path: Path) -> str:
        """解析纯文本文件"""
        with open(path, "r", encoding="utf-8") as f:
            return f.read()
    
    def clean_text(self, text: str) -> str:
        """
        清理文本
        
        - 去除多余空白
        - 去除特殊字符
        - 规范化换行
        """
        text = re.sub(r"\s+", " ", text)
        text = text.strip()
        text = re.sub(r"\n{3,}", "\n\n", text)
        return text
